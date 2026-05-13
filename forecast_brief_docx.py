"""
Forecast Approach & Results Brief — DOCX export.

Mirrors forecast_brief.py content in Word format. Hero chart is the
anonymized eventaxis bar chart; page 3 includes scatter, class, and
accuracy charts.
"""
import os
from datetime import date
import pandas as pd
from docx import Document
from docx.shared import Inches, Pt, RGBColor
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.enum.table import WD_ALIGN_VERTICAL
from docx.oxml.ns import qn
from docx.oxml import OxmlElement

WORKING_DIR = "/Users/antho/Documents/WPI-MW"
os.chdir(WORKING_DIR)

NAVY   = RGBColor(0x1A, 0x3A, 0x5C)
ORANGE = RGBColor(0xE8, 0x92, 0x2A)
TEAL   = RGBColor(0x2A, 0x9E, 0xA0)
INK    = RGBColor(0x1A, 0x23, 0x30)
INK_LT = RGBColor(0x5A, 0x6A, 0x7A)


HOLDOUT = [
    ("2022\u201323", 20, 27.3, 31.7),
    ("2023\u201324", 21, 28.3,  6.3),
    ("2024\u201325", 23, 23.1, -5.9),
]
HOLDOUT_TOTAL = (64, 26.3, 9.8)


def live_stats():
    df = pd.read_excel("Forecast_2526_FullSeason.xlsx")
    done = df[df["Status"] == "Completed"].copy()
    wape = (done["Pred_Adj"] - done["Actual"]).abs().sum() / done["Actual"].sum() * 100
    bias = ((done["Pred_Adj"] - done["Actual"]) / done["Actual"]).mean() * 100
    return wape, bias, len(done), len(df) - len(done)


def _set_cell_shading(cell, hex_color):
    tc_pr = cell._tc.get_or_add_tcPr()
    shd = OxmlElement("w:shd")
    shd.set(qn("w:fill"), hex_color)
    tc_pr.append(shd)


def _add_run(p, text, *, bold=False, italic=False, size=None, color=None, font="Open Sans"):
    r = p.add_run(text)
    r.font.name = font
    if size is not None:
        r.font.size = Pt(size)
    if color is not None:
        r.font.color.rgb = color
    r.bold = bold
    r.italic = italic
    return r


def _h1(doc, text):
    p = doc.add_paragraph()
    p.paragraph_format.space_before = Pt(0)
    p.paragraph_format.space_after = Pt(2)
    _add_run(p, text, bold=True, size=20, color=NAVY, font="Montserrat")


def _subtitle(doc, text):
    p = doc.add_paragraph()
    p.paragraph_format.space_before = Pt(0)
    p.paragraph_format.space_after = Pt(10)
    _add_run(p, text, size=10, color=INK_LT)
    pPr = p._p.get_or_add_pPr()
    pBdr = OxmlElement("w:pBdr")
    bottom = OxmlElement("w:bottom")
    bottom.set(qn("w:val"), "single")
    bottom.set(qn("w:sz"), "4")
    bottom.set(qn("w:space"), "4")
    bottom.set(qn("w:color"), "E8EDF2")
    pBdr.append(bottom)
    pPr.append(pBdr)


def _h2(doc, text):
    p = doc.add_paragraph()
    p.paragraph_format.space_before = Pt(10)
    p.paragraph_format.space_after = Pt(3)
    _add_run(p, text.upper(), bold=True, size=11, color=NAVY, font="Montserrat")


def _para(doc, runs, *, space_after=6, size=10.5, italic_caption=False):
    """runs: list of (text, {bold, italic}) or plain str."""
    p = doc.add_paragraph()
    p.paragraph_format.space_before = Pt(0)
    p.paragraph_format.space_after = Pt(space_after)
    for r in runs:
        if isinstance(r, str):
            _add_run(p, r, size=size, color=INK, italic=italic_caption)
        else:
            text, kw = r
            _add_run(p, text, size=size, color=kw.get("color", INK),
                     bold=kw.get("bold", False), italic=kw.get("italic", italic_caption))
    return p


def _bullet(doc, runs, *, size=10.5):
    p = doc.add_paragraph(style="List Bullet")
    p.paragraph_format.space_before = Pt(0)
    p.paragraph_format.space_after = Pt(3)
    for r in runs:
        if isinstance(r, str):
            _add_run(p, r, size=size, color=INK)
        else:
            text, kw = r
            _add_run(p, text, size=size, color=kw.get("color", INK),
                     bold=kw.get("bold", False), italic=kw.get("italic", False))
    return p


def _caption(doc, text):
    p = doc.add_paragraph()
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    p.paragraph_format.space_before = Pt(0)
    p.paragraph_format.space_after = Pt(6)
    _add_run(p, text, size=8.5, color=INK_LT, italic=True)


def _image(doc, path, *, width_in=6.6):
    p = doc.add_paragraph()
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    p.paragraph_format.space_before = Pt(2)
    p.paragraph_format.space_after = Pt(0)
    p.add_run().add_picture(path, width=Inches(width_in))


def _callout(doc, runs):
    """Single-cell shaded box used for the WAPE/Bias definition."""
    tbl = doc.add_table(rows=1, cols=1)
    tbl.autofit = True
    cell = tbl.rows[0].cells[0]
    _set_cell_shading(cell, "F8FAFC")
    p = cell.paragraphs[0]
    p.paragraph_format.space_before = Pt(2)
    p.paragraph_format.space_after = Pt(2)
    for r in runs:
        if isinstance(r, str):
            _add_run(p, r, size=9.5, color=INK)
        else:
            text, kw = r
            _add_run(p, text, size=9.5, color=kw.get("color", INK),
                     bold=kw.get("bold", False), italic=kw.get("italic", False))
    # Left border accent
    tcPr = cell._tc.get_or_add_tcPr()
    tcBorders = OxmlElement("w:tcBorders")
    left = OxmlElement("w:left")
    left.set(qn("w:val"), "single")
    left.set(qn("w:sz"), "18")
    left.set(qn("w:color"), "2A9EA0")
    tcBorders.append(left)
    tcPr.append(tcBorders)


def _results_table(doc, wape, bias, n_done, n_upc):
    headers = ["Season", "n", "WAPE", "Bias"]
    body = [
        (s, str(n), f"{w:.1f}%", f"{b:+.1f}%") for s, n, w, b in HOLDOUT
    ]
    n_tot, w_tot, b_tot = HOLDOUT_TOTAL
    body.append(("3-season total", str(n_tot), f"{w_tot:.1f}%", f"{b_tot:+.1f}%"))
    body.append((
        "2025\u201326 live (in-progress)",
        f"{n_done} of {n_done + n_upc}",
        f"{wape:.1f}%",
        f"{bias:+.1f}%",
    ))

    tbl = doc.add_table(rows=1 + len(body), cols=4)
    tbl.autofit = False
    widths = [Inches(2.4), Inches(1.0), Inches(1.0), Inches(1.0)]
    for i, w in enumerate(widths):
        for row in tbl.rows:
            row.cells[i].width = w

    # Header
    for j, h in enumerate(headers):
        c = tbl.rows[0].cells[j]
        p = c.paragraphs[0]
        if j > 0:
            p.alignment = WD_ALIGN_PARAGRAPH.RIGHT
        _add_run(p, h, bold=True, size=9.5, color=INK_LT)

    for i, row in enumerate(body, start=1):
        is_total = i == len(body) - 1
        is_live = i == len(body)
        for j, val in enumerate(row):
            c = tbl.rows[i].cells[j]
            p = c.paragraphs[0]
            if j > 0:
                p.alignment = WD_ALIGN_PARAGRAPH.RIGHT
            color = NAVY if (is_total or is_live) else INK
            _add_run(p, val, size=9.5, color=color,
                     bold=(is_total or is_live))
            if is_live:
                _set_cell_shading(c, "FFF6E8")


def build_docx():
    wape, bias, n_done, n_upc = live_stats()
    today = f"{date.today():%B %Y}"

    doc = Document()
    # Letter, narrow margins
    for section in doc.sections:
        section.page_height = Inches(11)
        section.page_width  = Inches(8.5)
        section.top_margin    = Inches(0.55)
        section.bottom_margin = Inches(0.55)
        section.left_margin   = Inches(0.65)
        section.right_margin  = Inches(0.65)

    # Default font
    style = doc.styles["Normal"]
    style.font.name = "Open Sans"
    style.font.size = Pt(10.5)

    # ===== Page 1 =====
    _h1(doc, "How many seats will this fill?")
    _subtitle(doc, f"A working attendance forecast for a regional performing-arts season  ·  {today}")

    _para(doc, [
        "A regional presenter signs next season's contracts 12\u201318 months out: "
        "venues, artists, fees, marketing budget. Each event is a bet on a single "
        "number \u2014 how many people will show up. Miss high and you're staring at "
        "empty seats and a marketing post-mortem; miss low and you've under-resourced "
        "the show that needed the most help. Most arts organizations make those bets "
        "on instinct. This brief is about a forecasting model that doesn't.",
    ])

    _para(doc, [
        "Built and running in production at a regional presenter, the model produces "
        "per-event attendance estimates at the start of the planning cycle \u2014 before "
        "a single ticket goes on sale. Through ",
        (f"{n_done} of {n_done + n_upc}", {"bold": True}),
        " completed events of the 2025\u201326 season, those pre-season predictions are "
        "tracking within ",
        (f"{wape:.0f}%", {"bold": True}),
        " of actual attendance, with effectively no directional drift (bias ",
        (f"{bias:+.0f}%", {"bold": True}),
        ").",
    ])

    _h2(doc, "The 2025\u201326 scorecard")
    _image(doc, "forecast_2526_bar_chart_eventaxis_anon.png", width_in=7.0)
    _caption(doc, "Predicted (navy) vs. actual (orange) attendance per event. Lighter "
                  "tones are comps. Anonymized \u2014 identifiers and scale removed.")

    _para(doc, [
        "The bars line up. That's the headline. Some events came in slightly under, "
        "some slightly over, none catastrophically. The largest misses sit on the "
        "model's known weak edges \u2014 programs without clean comparable history, or "
        "recitals where the artist-signal layer hasn't yet seen enough sample to "
        "anchor. The closest calls were headliners with strong bucket precedent and a "
        "recurring-series tag \u2014 a vote for the model's strength on its core repertoire.",
    ])

    _h2(doc, "What it changes in practice")
    _para(doc, [
        "Going into a season with a defensible per-event attendance number changes "
        "three conversations. ",
        ("Budgeting", {"bold": True}),
        " stops being aggregate hand-waving \u2014 each event has a forecast and the "
        "season's revenue picture is the sum of those, auditable line by line. ",
        ("Season planning", {"bold": True}),
        " gets a common scale: a proposed booking arrives with a draw estimate in the "
        "same units as everything else on the calendar, so \u201cdo we have enough "
        "capacity at this draw tier?\u201d becomes a question with a number, not a vibe. "
        "And ",
        ("marketing", {"bold": True}),
        " prioritizes earlier \u2014 events forecast under capacity surface as candidates "
        "for paid promotion or partnership outreach months before sale data would "
        "confirm the gap.",
    ])

    doc.add_page_break()

    # ===== Page 2 =====
    _h1(doc, "Inside the model")
    _subtitle(doc, "Why the standard tools don't work, and what does.")

    _para(doc, [
        "The forecasting problem in performing arts isn't usually too much data \u2014 "
        "it's too little of the right kind. A typical season is 25\u201335 events spanning "
        "headliner orchestras, chamber recitals, jazz combos, choral programs, "
        "education events, and free pay-what-you-want concerts. A cell defined by the "
        "natural cuts \u2014 class \u00d7 venue \u00d7 subgenre \u00d7 line of business \u2014 might "
        "contain a single prior observation, or none at all. Rule-of-thumb averages "
        "(\u201ca recital in this hall draws roughly 500\u201d) collapse the signal that "
        "matters most: artist stature, repeat-series momentum, pricing format, venue tier.",
    ])

    _para(doc, [
        "The model handles this through a ",
        ("five-level fallback hierarchy", {"bold": True}),
        ". Each event finds its prediction at the closest comparable cell available, "
        "falling back to progressively broader pools until it lands on a stable mean. ",
        ("Empirical-Bayes shrinkage", {"bold": True}),
        " at each level pulls thin buckets toward their next-coarser fallback \u2014 "
        "niche low buckets are protected, but thin buckets sitting unrealistically "
        "high get pulled back. A ",
        ("recurring-series prior", {"bold": True}),
        " fires before the hierarchy whenever a series has \u22652 prior observations, "
        "capturing the momentum of annual programs and returning headliners.",
    ])

    _para(doc, [
        "Two layers fill in where the bucket structure runs thin. PWYW events receive "
        "a shrunken multiplicative lift from historical PWYW sample. Venue-tier "
        "pooling (Marquee through Intimate) provides a fallback when subgenre \u00d7 venue "
        "cells are empty. Finally, an ",
        ("artist-popularity layer", {"bold": True}),
        " absorbs Wikipedia, Last.fm, and Deezer signals through an informed-Bayesian "
        "regression on log(actual / bucket prior). A signal-strength gate fires the "
        "adjustment only when genre-fit and signal thresholds are met \u2014 preventing "
        "global-pop signals from distorting world, folk, and Americana events where "
        "they don't apply.",
    ])

    _para(doc, [
        "Evaluation is ",
        ("temporal-holdout", {"bold": True}),
        ": each test season is predicted using only data strictly prior to it, with "
        "pandemic-era events down-weighted. There's no way for future information to "
        "leak into a past forecast \u2014 which is what makes the multi-season numbers "
        "below defensible.",
    ])

    _h2(doc, "Three seasons of evidence")
    _results_table(doc, wape, bias, n_done, n_upc)

    p = doc.add_paragraph()
    p.paragraph_format.space_before = Pt(2)
    p.paragraph_format.space_after = Pt(6)
    _add_run(p,
        "Error has compressed across seasons as the model stabilized. The 2022\u201323 "
        "positive bias reflects the early model over-predicting against still-recovering "
        "pandemic-era demand; the season-weighting scheme and shrinkage layers since "
        "then have largely closed that gap.",
        size=9, color=INK_LT)

    _callout(doc, [
        ("WAPE", {"bold": True, "color": NAVY}),
        " (Weighted Absolute Percentage Error) \u2014 total miss divided by total actual "
        "attendance. Volume-weighted so misses on high-draw events count more than "
        "misses on low-draw events. ",
        ("Bias", {"bold": True, "color": NAVY}),
        " \u2014 average signed error; near zero indicates no systematic over- or "
        "under-prediction.",
    ])

    _h2(doc, "What's coming")
    _bullet(doc, [
        ("In-season blending. ", {"bold": True}),
        "Combine the pre-season forecast with running sale counts via Kaplan\u2013Meier "
        "survival curves of ticket-purchase timing by event class \u2014 a forecast that "
        "sharpens as the show approaches.",
    ])
    _bullet(doc, [
        ("Sales-pace and pricing. ", {"bold": True}),
        "Use the same temporal curves to test where late-sale discounting erodes "
        "revenue versus fills otherwise-empty houses, and where higher-tier price "
        "headroom exists.",
    ])
    _bullet(doc, [
        ("Class-specific slopes ", {"bold": True}),
        "in the artist-adjustment layer once each event-class bucket reaches "
        "~15\u201320 observations.",
    ])

    _h2(doc, "About this work")
    _para(doc, [
        "The model and supporting analyses were built by Nolichucky Associates for a "
        "regional performing-arts presenter. This brief is anonymized for sharing "
        "with academic, peer-analytics, and arts-organization audiences. For "
        "collaboration, comparison studies, or replication discussions: "
        "nolichuckyassociates.com.",
    ])

    doc.add_page_break()

    # ===== Page 3 =====
    _h1(doc, "Looking deeper")
    _subtitle(doc, "Calibration, error by event class, and accuracy trend over the past three seasons.")

    _h2(doc, "Calibration \u2014 predicted vs. actual, 2025\u201326")
    _image(doc, "forecast_portfolio_scatter_2526.png", width_in=4.6)
    _caption(doc, "Each point is a completed event. Points on the dashed diagonal "
                  "indicate perfect prediction; above the line means the event drew "
                  "more than forecast, below means less. Color reflects event class.")

    _h2(doc, "Where the model fits best \u2014 by event class, 2025\u201326")
    _image(doc, "forecast_portfolio_class_2526.png", width_in=6.5)
    _caption(doc, "Headliners and the new Prestige tier (specialist artists with "
                  "strong core-audience appeal) are predicting tightly. Standard "
                  "programming carries the largest share of variance \u2014 its breadth "
                  "across genre, venue, and repeat-status is also the broadest of any class.")

    _h2(doc, "Accuracy trend \u2014 past three seasons + live")
    _image(doc, "forecast_portfolio_accuracy.png", width_in=6.0)
    _caption(doc, "Per-event MAPE (bars) and average bias (line). Error has compressed "
                  "over three seasons as the season-weighting scheme, shrinkage layers, "
                  "and artist-popularity layer have accumulated. The live 2025\u201326 bar "
                  "reflects pre-Dec 2025 events; the full season is at WAPE 19% / "
                  "bias +3% (see Page 1).")

    out = "forecast_brief.docx"
    doc.save(out)
    print(f"  \u2713 {out}")
    return out


if __name__ == "__main__":
    build_docx()
