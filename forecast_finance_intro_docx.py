"""
Finance-side intro brief — DOCX, simple formatting.

For non-technical finance audience supporting 26-27 budget planning.
Frames the 25-26 hindcast as model validation, surfaces lessons learned,
and gives a forward look at 26-27. Uses the wide 2-column horizontal
charts for readability.

Note: this script reflects the user's edited working version. For prose
tweaks, edit the resulting docx directly — don't regenerate unless the
numbers, charts, or structure need to update.

Variants:
  python forecast_finance_intro_docx.py                    → forecast_finance_intro.docx              (named, for MW finance)
  python forecast_finance_intro_docx.py --anon             → forecast_finance_intro_anon.docx         (org-anonymized, percentages only)
  python forecast_finance_intro_docx.py --anon --scale     → forecast_finance_intro_anon_scaled.docx  (anon but chart $ scale visible)
"""
import os
import sys
import pandas as pd
from docx import Document
from docx.shared import Inches, Pt
from docx.enum.text import WD_ALIGN_PARAGRAPH

ANON  = "--anon" in sys.argv
SCALE = "--scale" in sys.argv
WORKING_DIR = "/Users/antho/Documents/WPI-MW"
os.chdir(WORKING_DIR)


def load_attendance():
    df = pd.read_excel("forecasting/Forecast_2526_FullSeason.xlsx", sheet_name="Detail")
    done = df[df["Status"] == "Completed"].copy()
    wape = (done["Pred_Adj"] - done["Actual"]).abs().sum() / done["Actual"].sum() * 100
    bias = ((done["Pred_Adj"] - done["Actual"]) / done["Actual"]).mean() * 100
    return {
        "n":      len(done),
        "actual": int(done["Actual"].sum()),
        "pred":   int(done["Pred_Adj"].sum()),
        "wape":   wape,
        "bias":   bias,
    }


def load_revenue():
    per_ev = pd.read_excel(
        "forecasting/Forecast_2526_Revenue_Hindcast.xlsx", sheet_name="Per-Event")
    per_ev = per_ev[per_ev["Event"].notna() & (per_ev["Event"] != "Total")].copy()
    pred = int(pd.to_numeric(per_ev["Pred Revenue"], errors="coerce").sum())
    act  = int(pd.to_numeric(per_ev["Actual Revenue"], errors="coerce").sum())
    return {
        "pred":    pred,
        "actual":  act,
        "gap":     pred - act,
        "gap_pct": (pred - act) / act * 100,
    }


def _caption(doc, text):
    cap = doc.add_paragraph()
    cap.alignment = WD_ALIGN_PARAGRAPH.CENTER
    run = cap.add_run(text)
    run.italic = True
    run.font.size = Pt(9)


def _image(doc, path, width_in=7.0):
    p = doc.add_paragraph()
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    p.add_run().add_picture(path, width=Inches(width_in))


def build_docx():
    att = load_attendance()
    rev = load_revenue()

    doc = Document()

    if ANON:
        title_text = "Performing Arts Pre-Season Forecasting"
        org_label  = "The presenter"
        if SCALE:
            att_chart = "forecasting/forecast_2526_bar_chart_wide_anon.png"
            rev_chart = "forecasting/forecast_revenue_chart_anon_scaled_horizontal.png"
        else:
            att_chart = "forecasting/forecast_2526_bar_chart_wide_anon.png"
            rev_chart = "forecasting/forecast_revenue_chart_anon_horizontal.png"
        marquee_evidence = ("The most prestigious orchestra booking this "
                            "season was consistent with the prior "
                            "Yo-Yo Ma / Itzhak Perlman class of bookings, "
                            "which we've used to anchor the Marquee tier "
                            "going forward")
    else:
        title_text = "Music Worcester - Pre-Season Forecasting"
        org_label  = "MW"
        att_chart  = "forecasting/forecast_2526_bar_chart_wide.png"
        rev_chart  = "forecasting/forecast_revenue_chart_horizontal.png"
        marquee_evidence = ("ONF/Trifonov was consistent with the Yo-Yo Ma / "
                            "Itzhak Perlman class of past bookings, which "
                            "we've used to anchor the Marquee tier going "
                            "forward")

    # ── Title + intro ──
    doc.add_heading(title_text, level=1)

    doc.add_paragraph(
        f"{org_label} needs an attendance and revenue estimate for every "
        f"event in the upcoming season to support budgeting and planning. "
        f"To help with this, we've developed a forecast model that produces "
        f"a forecast based on 13 years of historical ticket sales "
        f"(favoring more recent years) and an event manifest. The manifest "
        f"is a list of every event with its venue, capacity, programming "
        f"type, pricing tier, and so on."
    )
    doc.add_paragraph(
        f"The forecast is completely automated and runs in seconds, so "
        f"that mid-season updates to zero in on upcoming events with new "
        f"sales data are very easy."
    )
    doc.add_paragraph(
        f"To test it, we've “hindcasted” the 25-26 season by "
        f"blinding the model to 25-26 sales, letting the model predict "
        f"the 25-26 season and comparing the forecast to the actuals. "
        f"This helps us refine the model and gives us confidence that the "
        f"26-27 forecast is useful for budget planning. Here are the "
        f"results and some lessons learned."
    )

    # ── Attendance hindcast ──
    doc.add_heading("Attendance - 25-26 hindcast results", level=2)
    if ANON:
        att_line = (
            f"Net forecast landed within "
            f"{abs((att['pred']-att['actual'])/att['actual'])*100:.2f}% of "
            f"actual attendance. Over {att['n']} events of the 25-26 season, "
            f"the per-event spread (WAPE) averaged {att['wape']:.0f}% and "
            f"overall bias was {att['bias']:.0f}%. Individual misses go both "
            f"ways and cancel when added together. And the results are "
            f"pretty solid."
        )
    else:
        att_line = (
            f"Net forecast: {att['pred']:,} vs. actual {att['actual']:,} "
            f"(off by {att['pred']-att['actual']:+,} of 14,000, within "
            f"{abs((att['pred']-att['actual'])/att['actual'])*100:.2f}%). "
            f"Over {att['n']} events of the 25-26 season, the per-event "
            f"spread (WAPE) averaged {att['wape']:.0f}% and overall bias "
            f"was {att['bias']:.0f}%. Individual misses go both ways and "
            f"cancel when added together. And the results are pretty solid."
        )
    doc.add_paragraph(att_line)
    doc.add_paragraph(
        f"The model works by averaging past similar events, where "
        f"similarity is measured along several dimensions: event class "
        f"(Marquee, Headliner, Standard, etc.), venue, genre, "
        f"recurring-series (Messiah, DTH, Cantatathon, etc.), and the "
        f"artist's online popularity. Closer matches and more recent "
        f"seasons carry more weight."
    )
    _image(doc, att_chart)
    _caption(doc, "Predicted (navy) vs. actual (orange) attendance by event."
                   + (" Scale removed." if (ANON and not SCALE) else ""))

    # ── Revenue hindcast ──
    doc.add_heading("Revenue - 25-26 hindcast results", level=2)
    doc.add_paragraph(
        f"Revenue forecasting is based on the attendance forecast, "
        f"historical ticket prices, and our projections of ticket price "
        f"increases for the various pricing tiers."
    )
    if ANON and not SCALE:
        doc.add_paragraph(
            f"Revenue forecast landed within {abs(rev['gap_pct']):.1f}% of "
            f"actual at the season level."
        )
    else:
        doc.add_paragraph(
            f"Revenue forecast: ${rev['pred']:,} vs. actual ${rev['actual']:,} "
            f"(within {abs(rev['gap_pct']):.1f}% at the season level)."
        )
    doc.add_paragraph(
        f"Event revenue = predicted paid attendance from the forecast × "
        f"historical average ticket prices × pricing tier uplift."
    )
    p = doc.add_paragraph(style="List Bullet")
    p.add_run(
        "The historical ticket prices come from similar past events and "
        "we add % uplift adjustments for each tier based on my recollection "
        "of pre-season pricing discussions for TCB, ONF, etc."
    )
    p = doc.add_paragraph(style="List Bullet")
    p.add_run(
        "The price tier uplifts are: Marquee +40%, Headliner/Prestige +15%, "
        "Standard +10%, Chorus/Mission/AiR 0%, TCB Main 0%, TCB Side −25% "
        "(the negative uplift reflects our audience-building accessible "
        "pricing for festival side events)."
    )
    doc.add_paragraph(
        f"And the revenue forecast is a work in progress, so expect further "
        f"refinements."
    )
    _image(doc, rev_chart)
    _caption(doc, "Predicted (navy) vs. actual (orange) revenue by event."
                   + (" Scale removed." if (ANON and not SCALE) else ""))

    # ── Key lessons (retrospective) ──
    doc.add_heading("Key lessons from 25-26 hindcast", level=2)

    p = doc.add_paragraph(style="List Bullet")
    p.add_run("Top-tier Marquee bookings have pricing headroom. ").bold = True
    p.add_run(
        f"{marquee_evidence}. Future top marquee bookings may still "
        f"surprise on the upside."
    )

    p = doc.add_paragraph(style="List Bullet")
    p.add_run("Volume and price both respond on Marquee bookings. ").bold = True
    p.add_run(
        "ONF/Trifonov was forecast at 823 paid at $76, drew 1,089 at $97 — "
        "both attendance and price came in well above the model. The same "
        "handful of top-tier events tend to drive both attendance and "
        "revenue surprises, and they account for most of the aggregate "
        "revenue gap."
    )

    p = doc.add_paragraph(style="List Bullet")
    p.add_run("Selective above-tier pricing on individual events worked. ").bold = True
    p.add_run(
        "When MW prices an event above its tier's typical range, the "
        "model can't anticipate that. This season Sebastians, Savall, and "
        "Dinnerstein Recital were priced 20–30% above what the tier model "
        "expected, and audiences absorbed the higher prices. For 26-27, "
        "when MW plans above-tier pricing on specific events, we can add "
        "a manual revenue adjustment on top of the model for those line "
        "items. We can also look at refining/expanding the tier structure."
    )

    p = doc.add_paragraph(style="List Bullet")
    p.add_run("TCB and Chorus events are priced to drive audiences. ").bold = True
    p.add_run(
        "TCB festival side events (organ recitals, cantatas) consistently "
        "run $20–30 tickets; Worcester Chorus events similarly. These are "
        "intentional decisions, not underpricing misses. The model encodes "
        "this with 0% uplift on Chorus / Mission / AiR / TCB Main and "
        "−25% on TCB Side. These tiers shouldn't be read as opportunities "
        "to raise prices."
    )

    # ── Looking ahead to 26-27 ──
    doc.add_heading("Looking ahead to 26-27 forecast", level=2)

    p = doc.add_paragraph(style="List Bullet")
    p.add_run("Season and event-level forecasts for 26-27 should be reliable for budget-level planning. ").bold = True
    p.add_run(
        "Build in ±15–20% per-event tolerance for marketing or capacity decisions."
    )

    p = doc.add_paragraph(style="List Bullet")
    p.add_run("Identify Marquee-class bookings pre-season and tag them. ").bold = True
    p.add_run(
        "The Marquee tier (+40% uplift) is anchored to historical Yo-Yo "
        "Ma / Perlman / Vengerov bookings. When a 26-27 booking fits that "
        "class (1,000+ expected paid at $80+ pricing in Mechanics), we "
        "should tag it as Marquee in the manifest to get the right "
        "revenue projection."
    )

    p = doc.add_paragraph(style="List Bullet")
    p.add_run("The new Artist in Residence next season will likely outperform the historical AiR events. ").bold = True
    p.add_run(
        "The model is anchored to prior AiR events, which have been "
        "community-priced and mission-focused. Everett McCorvey will be "
        "more popular, so the model will likely under-predict both "
        "attendance and revenue for his residency. A manual uplift on "
        "the AiR forecast may be appropriate once the booking details "
        "and pricing posture are set."
    )

    p = doc.add_paragraph(style="List Bullet")
    p.add_run("TCB models should improve. ").bold = True
    p.add_run(
        "The 25-26 hindcast only had 1 year of complete TCB events. With "
        "one more year of data, we can expect the TCB and festival event "
        "forecasts to be more refined."
    )

    # ── Source files ──
    doc.add_heading("Source files:", level=2)

    p = doc.add_paragraph(style="List Bullet")
    p.add_run("Forecast_2526_FullSeason.xlsx (attendance),")

    p = doc.add_paragraph(style="List Bullet")
    p.add_run("Forecast_2526_Revenue_Hindcast.xlsx (revenue).")

    out = ("forecasting/forecast_finance_intro_anon_scaled.docx" if (ANON and SCALE)
           else "forecasting/forecast_finance_intro_anon.docx" if ANON
           else "forecasting/forecast_finance_intro.docx")
    doc.save(out)
    print(f"  ✓ {out}")


if __name__ == "__main__":
    build_docx()
